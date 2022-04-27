# Import docx NOT python-docx
import docx
from docx.shared import Pt

# Create an instance of a word document
import Utility

doc = docx.Document()

# Adding paragraph with Increased font size
doc.add_heading('\tIncrease Font Size', 3)

para = doc.add_paragraph().add_run('Murali')
# Increasing size of the font
para.font.size = Pt(30)

# Adding paragrap
doc.save(Utility.projectDirectory()+"8.SettingTheFontSize.docx")