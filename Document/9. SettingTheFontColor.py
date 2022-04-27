#To apply a font colour to the text you have to first create a paragraph object then you have to use add_run() method to add content. You can directly use add_paragraph() method to add paragraph but if you want to apply a font colour to a text you have to use add_run() as all the block-level formatting is done by using add_paragraph() method while all the character-level formatting is done by using add_run()

import docx
from docx.shared import RGBColor

# Create an instance of a word document
import Utility

doc = docx.Document()

# Add a Title to the document
doc.add_heading('Setting the color', 0)

para = doc.add_paragraph().add_run('Murali')

# Adding forest green colour to the text
# RGBColor(R, G, B)
para.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)

# Now save the document to a location
doc.save(Utility.projectDirectory()+"9.SettingTheFontColor.docx")