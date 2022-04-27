#Author - Murali.r
#Date - 27th April 2022
#GitHub ID -
from docx import Document

import Utility

document = Document()

# adding pic in header
section = document.sections[0]

header = section.header

paragraph = header.paragraphs[0]

run = paragraph.add_run()

run.add_picture(Utility.projectDirectoryPath()+"\\Dependency\\Google.png")

document.save(Utility.projectDirectory()+"AddPictureToTheDocument.docx")