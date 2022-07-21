f = 'E:\\iphone\\IQ.docx'

import docx

import docx

doc = docx.Document(f)

print(doc.paragraphs[0].text)

all_paras = doc.paragraphs

print(len(all_paras))

for para in all_paras:
    print(para.text)
