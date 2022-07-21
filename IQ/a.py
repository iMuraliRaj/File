# Import docx NOT python-docx
import shutil
from docx.shared import Pt, RGBColor
import docx
from docx.shared import Inches, Cm
doc = docx.Document()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.06)
    section.right_margin = Cm(2.54)

para = doc.add_paragraph().add_run('\n\n\n\n\n         Qualis LIMS (ReactJS) 1.0')
# Increasing size of the font
para.font.size = Pt(28)
para.bold=True
para.alignment=1
para.font.name = 'Arial'

para = doc.add_paragraph().add_run('          Installation Qualification\n')
# Increasing size of the font
para.font.size = Pt(28)
para.bold=True
para.alignment=1
para.font.name = 'Arial'


para = doc.add_paragraph().add_run('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n     This information is believed to be complete and accurate as of publication, and is subject to change without notice.')
# Increasing size of the font
para.font.size = Pt(9)
para.alignment=1
para.font.name = 'Arial'

para = doc.add_paragraph().add_run('                                  Copyright © 2000—2010 Agaram Technologies. All rights reserved.\n')
# Increasing size of the font
para.font.size = Pt(9)
para.font.name = 'Arial'

doc.add_page_break()

# Adding a paragraph
contentsHeading=doc.add_paragraph().add_run('Contents')
contentsHeading.font.size = Pt(14)
contentsHeading.bold=True
contentsHeading.font.color.rgb = RGBColor(6, 4, 255)
contentsHeading.font.name = 'Cambria'

section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
a = """-----------------------------------------------------------------------------------------------------------------------------------
The objective of this protocol is to verify the protocol,
The objective of this protocol is to verify the protocol,
The objective of this protocol is to verify the protocol
The objective of this protocol is to verify the protocol"""
footer_para.text =a

source='D:\\iMuraliRaj\\GitHub\\File\\IQ\\Installation Qulification.docx'
doc.save(source)
dest = shutil.copyfile(source, "C:\\Users\\Murali.R\\Desktop\\IQ.docx")