# Import docx NOT python-docx
import shutil
from docx.shared import Pt, RGBColor
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


from docx.shared import Inches, Cm
doc = docx.Document()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.06)
    section.right_margin = Cm(2.54)

para = doc.add_paragraph().add_run('\n\n\n\n\n         Qualis LIMS (ReactJS) 1.0')
# Increasing size of the font
para.font.size = Pt(28)
para.bold=True
para.alignment=1
para.font.name = 'Arial'

para = doc.add_paragraph().add_run('          Installation Qualification\n')
# Increasing size of the font
para.font.size = Pt(28)
para.bold=True
para.alignment=1
para.font.name = 'Arial'


para = doc.add_paragraph().add_run('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n     This information is believed to be complete and accurate as of publication, and is subject to change without notice.')
# Increasing size of the font
para.font.size = Pt(9)
para.alignment=1
para.font.name = 'Arial'

para = doc.add_paragraph().add_run('                                  Copyright © 2000—2010 Agaram Technologies. All rights reserved.\n')
# Increasing size of the font
para.font.size = Pt(9)
para.font.name = 'Arial'

doc.add_page_break()

# Adding a paragraph
contentsHeading=doc.add_paragraph().add_run('Contents')
contentsHeading.font.size = Pt(14)
contentsHeading.bold=True
contentsHeading.font.color.rgb = RGBColor(6, 4, 255)
contentsHeading.font.name = 'Cambria'

purpose=doc.add_paragraph('  Purpose\n',  style='List Number')
doc.add_paragraph('  Purpose',  style='List Number 2')
doc.add_paragraph('  Scope\n',  style='List Number')
doc.add_paragraph('  Validation Methodology\n',  style='List Number')
doc.add_paragraph('  Acronyms\n',  style='List Number')
doc.add_paragraph('  System Description\n',  style='List Number')
doc.add_paragraph('  Introduction\n',  style='List Number 2')
doc.add_paragraph('  Work Flow\n',  style='List Number 2')
doc.add_paragraph('  Responsibilities\n',  style='List Number')
doc.add_paragraph('  Test Plan\n',  style='List Number')
doc.add_paragraph('  Prerequisites Review\n',  style='List Number 2')
doc.add_paragraph('  Computer System Specification and Power Supply\n',  style='List Number 2')
doc.add_paragraph('  Software Review\n',  style='List Number 2')
doc.add_paragraph('  Qualis Installation Verification\n',  style='List Number 2')
doc.add_paragraph('  Verifying components of the Qualis\n',  style='List Number 2')
doc.add_paragraph('  Qualification Support Environment screen shots\n',  style='List Number')
doc.add_paragraph('  Deficiency and Change Request Log\n',  style='List Number')
doc.add_paragraph('  Document Approval',  style='List Number')

doc.add_page_break()
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]

# Adding the left zoned footer
header_para.text ="\tINSTALLATION QUALIFICATION"

table = header.add_table(1, 3, Inches(10))
row = table.add_row().cells
row[0].text = "INSTALL"


def headerParagraph(content):
    # Adding a paragraph
    contentsHeading = doc.add_paragraph().add_run(content)
    contentsHeading.font.size = Pt(14)
    contentsHeading.bold = True
    contentsHeading.font.color.rgb = RGBColor(6, 4, 255)
    contentsHeading.font.name = 'Cambria'

headerParagraph('1. Purpose')


contentsHeading=doc.add_paragraph().add_run('          The objective of this protocol is to verify the installation of the Qualis LIMS program and to \n          verify that it performs its intended function in a consistent and reproducible manner. Specific \n           test objectives and acceptance criteria are defined. ')
contentsHeading.font.size = Pt(11)
contentsHeading.font.name = 'Arial'

headerParagraph('1.1.   Installation Qualification')

contentsHeading=doc.add_paragraph().add_run('          The Installation Qualification consists of tests to ensure that the components of the Qualis \n          LIMS program have been properly installed. Specific test objectives and acceptance criteria \n           are defined.. ')
contentsHeading.font.size = Pt(11)
contentsHeading.font.name = 'Arial'

headerParagraph('2.  Scope')

scopeContent="""        This protocol specifies the Installation Qualification tests that are to be applied to the Qualis LIMS program, which will be used in a cGMP/cGLP environment, for the purpose of verifying proper installation. The Qualis LIMS is installed to manage Laboratory workflows and data generated in process over the LAN to a database. Within the scope of this protocol are the activities and tasks that must be performed throughout the qualification process.
Changes made to the software and/or computer systems after validation is completed must be conducted under the appropriate system change control procedure. Parts of the protocol may be used for reconfirmation or revalidation, if necessary."""

contentsHeading=doc.add_paragraph().add_run(scopeContent)
contentsHeading.font.size = Pt(11)
contentsHeading.font.name = 'Arial'

def content(content):
    contentsHeading=doc.add_paragraph().add_run(content)
    contentsHeading.font.size = Pt(11)
    contentsHeading.font.name = 'Arial'

headerParagraph("3. Validation Methodology")

validationContent="""Tests within this protocol have been designed to verify that all-important elements of the Qualis LIMS program installation adhere to the requirements set forth by the manufacturer, when successfully completed. For each installation test in this protocol, a cover sheet is provided to define the test objective, procedure, and acceptance criteria. Following each cover sheet are the data sheets. These data sheets outline the information that must be verified and/or documented, as well as provide space for recording additional information. The information recorded on the test data sheets combined with referenced supporting test documentation provides a method whereby adherence to the test acceptance criteria can be verified.
\nAny deviation from the specified acceptance criteria/expected results must be recorded on a deviation report."""

content(validationContent)

doc.add_page_break()
headerParagraph('4. Acronyms')

acronymsTable = doc.add_table(rows=24, cols=2 ,style="Table Grid")
acronymsTableRows = acronymsTable.rows[0].cells
acronymsTableRows[0].add_paragraph().add_run("Acronym/Initials\n").bold=True
acronymsTableRows[1].add_paragraph().add_run("Meaning\n").bold=True


row[0].add_paragraph().add_run("Deficiency number\n").bold=True
row[1].add_paragraph().add_run('Brief description').bold=True


doc.add_page_break()
headerParagraph('8. Qualification Support Environment screen shots')



def heading(content):
    contentsHeading = doc.add_paragraph().add_run(content)
    contentsHeading.font.size = Pt(11)
    contentsHeading.font.name = 'Calibri'
    contentsHeading.bold=True

heading("Tomcat")

doc.add_picture('D:\\BUILD\\IRSHA\\New folder\\Tomcat.png',width=Inches(7), height=Inches(3.5))

heading("ADS Services")

doc.add_picture('D:\\BUILD\\IRSHA\\New folder\\services.png',width=Inches(7), height=Inches(3.5))



heading("IIS")

doc.add_picture('D:\\BUILD\\IRSHA\\New folder\\IIS.png',width=Inches(7), height=Inches(3.5))

heading("Email")

doc.add_picture('D:\\BUILD\\IRSHA\\New folder\\Email.png',width=Inches(7), height=Inches(3.5))

doc.add_page_break()

heading("Application")

doc.add_picture('D:\\BUILD\\IRSHA\\New folder\\Application.png',width=Inches(7), height=Inches(3.5))

headerParagraph('\n\n9.  Deficiency and Change Request Log')

deficiency="""In the following Deficiency Log, record the deviation numbers for any created deviation reports. Include a brief description of the issue. When the deviation is resolved and approved, write the date resolved."""

doc.add_paragraph().add_run(deficiency)

table = doc.add_table(rows=7, cols=3 ,style="Table Grid")

shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#9b9b9b"/>'.format(nsdecls('w')))
table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#9b9b9b"/>'.format(nsdecls('w')))
table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

shading_elm_3 = parse_xml(r'<w:shd {} w:fill="#9b9b9b"/>'.format(nsdecls('w')))
table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_3)

for cell in table.columns[0].cells:
    cell.width = Inches(2)

for cell in table.columns[1].cells:
    cell.width = Inches(4)

for cell in table.rows:
    cell.height =  Cm(0.7)

table.rows[0].height = Cm(1)

row = table.rows[0].cells

row[0].add_paragraph().add_run("Deficiency number\n").bold=True
row[1].add_paragraph().add_run('Brief description').bold=True
row[2].add_paragraph().add_run('Date resolved').bold=True

doc.add_page_break()

headerParagraph('10.  Document Approval')


def content(content):
    contentsHeading = doc.add_paragraph().add_run(content)
    contentsHeading.font.size = Pt(10)
    contentsHeading.font.name = 'Calibri'

content("    The Operational Qualification document has been reviewed and approved.")



documentApprovalTable = doc.add_table(rows=6, cols=2 ,style="Table Grid")

for cell in documentApprovalTable.columns[0].cells:
    cell.width = Inches(7)

backround=""

shading_elm_4 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_4)

shading_elm_7 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_7)


shading_elm_5 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[2].cells[0]._tc.get_or_add_tcPr().append(shading_elm_5)

shading_elm_8 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[2].cells[1]._tc.get_or_add_tcPr().append(shading_elm_8)

shading_elm_6 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[4].cells[0]._tc.get_or_add_tcPr().append(shading_elm_6)

shading_elm_9 = parse_xml(r'<w:shd {} w:fill="#e7e7e7"/>'.format(nsdecls('w')))
documentApprovalTable.rows[4].cells[1]._tc.get_or_add_tcPr().append(shading_elm_9)


preparedByName="""Name:  Suganya. P

Printed Name: 
"""

preparedByDate="""Date: 21.04.2022

Implementation Engineer
Agaram Technologies """



row = documentApprovalTable.rows[0].cells
row[0].text="Prepared by"



row2=documentApprovalTable.rows[1].cells
row2[0].add_paragraph().add_run(preparedByName)
row2[1].add_paragraph().add_run(preparedByDate)

row = documentApprovalTable.rows[2].cells
row[0].text="Reviewed by"

reviewedBy="""Name: Satish P

Printed Name: """



reviewedByDate="""Date: 22.04.2022

Manager - Implementations
Agaram Technologies"""

row = documentApprovalTable.rows[3].cells
row[0].add_paragraph().add_run(reviewedBy)
row[1].add_paragraph().add_run(reviewedByDate)


approvedBy="""Name: 

Printed Name: 
"""


approvedByDate="""Date: 

LIMS Site Admin
NIBSC
"""


row = documentApprovalTable.rows[4].cells
row[0].text="Approved by"

row = documentApprovalTable.rows[5].cells
row[0].add_paragraph().add_run(approvedBy)
row[1].add_paragraph().add_run(approvedByDate)

source='D:\\iMuraliRaj\\GitHub\\File\\IQ\\Installation Qulification.docx'
doc.save(source)
dest = shutil.copyfile(source, "C:\\Users\\Murali.R\\Desktop\\IQ.docx")